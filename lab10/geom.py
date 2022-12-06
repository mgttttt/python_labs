import PySimpleGUI as sg
from docx import Document
from figures_calculations import rectangle, triangle, trapezoid
from figures_calculations.triangle import calc_len

sg.theme('DarkAmber')

layout1 = [[sg.Text("Select figure", )],
        [sg.Combo(["Triangle", "Rectangle", "Trapezoid"], default_value='Triangle'), sg.Button("OK")],
        [sg.Text("Input coordinates")],
        [sg.Text('x1'), sg.InputText(size=10), sg.Text('y1'), sg.InputText(size=10)],
        [sg.Text('x2'), sg.InputText(size=10), sg.Text('y2'), sg.InputText(size=10)],
        [sg.Text('x3'), sg.InputText(size=10), sg.Text('y3'), sg.InputText(size=10)],
        [sg.Text('x4', visible=False, key='x4'), sg.InputText(size=10, visible=False, key='x4i'), sg.Text('y4', visible=False, key='y4'), sg.InputText(size=10, visible=False, key='y4i')],
        [sg.Text("Result of calculations")],
        [sg.Text('Area'), sg.InputText(size=10, key='s'), sg.Text('Inscribed circle radius'), sg.InputText(size=10, key='r')],
        [sg.Text('Radius of the circumscribed circle'), sg.InputText(size=18, key='R')],
        [sg.Button('Make calculations'), sg.Button('Quit')]]


def write_docx(figure, coords, s, r, R):
    document = Document('geom.docx')
    document.add_paragraph(figure)
    document.add_paragraph('coordinates: ')
    for coord in coords:
        document.add_paragraph('(' + str(coord[0]) + ', ' + str(coord[1]) + ')')
    document.add_paragraph('Площадь:' + str(round(s, 5)))
    document.add_paragraph('Радиус вписанной окружности:' + str(round(r, 5)) if r else 'вписать окружность невозможно')
    document.add_paragraph('Радиус описанной окружности:' + str(round(R, 5)) if R else 'описать окружность невозможно')
    document.save('geom.docx')


def calc_area(values, figure):
    if figure == 'Rectangle' or figure == 'Trapezoid':
        coords = [(float(values[1]), float(values[2])),
                  (float(values[3]), float(values[4])),
                  (float(values[5]), float(values[6])),
                  (float(values['x4i']), float(values['y4i']))]
        coords = sorted(coords).copy()
        a, b, c, d, e = calc_len([coords[0], coords[1]]), calc_len([coords[0], coords[2]]), calc_len(
            [coords[1], coords[3]]), calc_len([coords[2], coords[3]]), calc_len([coords[0], coords[3]])
        if figure == 'Rectangle':
            s = rectangle.calc(a, b, c, d, e)
            r = rectangle.calc_r(a, b, c)
            R = rectangle.calc_R(a, b)
        else:
            s = trapezoid.calc(a, b, c, d, e)
            r = trapezoid.calc_r(a, b, c, d)
            R = trapezoid.calc_R(a, b, c, d, e)
        if not s:
            window.Element("s").Update('Error')
            window.Element("r").Update('Error')
            window.Element("R").Update('Error')
        else:
            window.Element("s").Update(str(round(s, 5)))
            window.Element("r").Update('can not find' if not r else str(round(r, 5)))
            window.Element("R").Update('can not find' if not R else str(round(R, 5)))
            write_docx(figure, coords, s, r, R)

    else:
        coords = [(float(values[1]), float(values[2])),
                  (float(values[3]), float(values[4])),
                  (float(values[5]), float(values[6]))]
        a, b, c = calc_len([coords[0], coords[1]]), calc_len([coords[0], coords[2]]), calc_len([coords[2], coords[1]])
        s = triangle.triangle_square(a, b, c)
        r = triangle.calc_r(s, a, b, c)
        R = triangle.calc_R(s, a, b, c)
        if s:
            window.Element("s").Update(str(round(s, 5)))
            window.Element("r").Update(str(round(r, 5)))
            window.Element("R").Update(str(round(R, 5)))
            write_docx(figure, coords, s, r, R)
        else:
            window.Element("s").Update('Error')
            window.Element("r").Update('Error')
            window.Element("R").Update('Error')


def visible_change(figure):
    if figure == 'Triangle':
        window.Element("x4").Update(visible=False)
        window.Element("x4i").Update(visible=False)
        window.Element("y4").Update(visible=False)
        window.Element("y4i").Update(visible=False)
    else:
        window.Element("x4").Update(visible=True)
        window.Element("x4i").Update(visible=True)
        window.Element("y4").Update(visible=True)
        window.Element("y4i").Update(visible=True)


window = sg.Window('Window Title', layout1)

while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Quit':
        break
    if event == 'OK':
        visible_change(values[0])
    if event == 'Make calculations':
        calc_area(values, values[0])
