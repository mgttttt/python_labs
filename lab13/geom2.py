from pyforms.basewidget import BaseWidget
from pyforms.controls import ControlText
from pyforms.controls import ControlButton
from pyforms.controls import ControlLabel
from pyforms.controls import ControlCombo
from pyforms import start_app
from docx import Document
import figures_calculations2.trapezoid2 as trapezoid2, figures_calculations2.rectangle2 as rectangle2, figures_calculations2.triangle2 as triangle2


class Geom_figures(BaseWidget):
    def __init__(self):
        super().__init__('Geometry figures')

        self.set_margin(10)

        self._text = ControlLabel('Select figure')
        self._combobox = ControlCombo()
        self._combobox.add_item('Triangle')
        self._combobox.add_item('Rectangle')
        self._combobox.add_item('Trapezoid')
        self._text2 = ControlLabel('Input coordinates')
        self._x1 = ControlText('x1')
        self._y1 = ControlText('y1')
        self._x2 = ControlText('x2')
        self._y2 = ControlText('y2')
        self._x3 = ControlText('x3')
        self._y3 = ControlText('y3')
        self._x4 = ControlText('x4')
        self._y4 = ControlText('y4')
        self._x4.readonly = True
        self._y4.readonly = True
        self._restext = ControlLabel('Result of calculations')
        self._area = ControlText('Area')
        self._r = ControlText('Inscribed circle radius')
        self._R = ControlText('Radius of the circumscribed circle')
        self._runbutton = ControlButton('Make calculations')
        self._runbutton.value = self.__runEvent
        self._combobox.activated_event = self.__itemChanged
        self._formset = [
            ('_text', '_combobox'),
            ('_text2'),
            ('_x1', '_y1'),
            ('_x2', '_y2'),
            ('_x3', '_y3'),
            ('_x4', '_y4'),
            ('_restext'),
            ('_area', '_r'),
            ('_R'),
            ('_runbutton')
        ]

    def __runEvent(self):
        if self._combobox.value == 'Triangle':
            coords = [(float(self._x1.value), float(self._y1.value)),
                      (float(self._x2.value), float(self._y2.value)),
                      (float(self._x3.value), float(self._y3.value))]
            coords = sorted(coords).copy()
            tr = triangle2.Triangle(coords)
            if tr.s:
                self._area.value = str(round(tr.s, 5))
                self._r.value = str(round(tr.r, 5))
                self._R.value = str(round(tr.R, 5))
                write_docx('Triangle', coords, tr.s, tr.r, tr.R)
                return
        else:
            coords = [(float(self._x1.value), float(self._y1.value)),
                      (float(self._x2.value), float(self._y2.value)),
                      (float(self._x3.value), float(self._y3.value)),
                      (float(self._x4.value), float(self._y4.value))]
            coords = sorted(coords).copy()
            if self._combobox.value == 'Rectangle':
                rec = rectangle2.Rectangle(coords)
                if rec.s:
                    self._area.value = str(round(rec.s, 5))
                    self._r.value = str(round(rec.r, 5))
                    self._R.value = str(round(rec.R, 5))
                    write_docx('Rectangle', coords, rec.s, rec.r, rec.R)
                    return
            else:
                trap = trapezoid2.Trapezoid(coords)
                if trap.s:
                    self._area.value = str(round(trap.s, 5))
                    self._r.value = str(round(trap.r, 5))
                    self._R.value = str(round(trap.R, 5))
                    write_docx('Trapezoid', coords, trap.s, trap.r, trap.R)
                    return
        self._area.value = 'Error'
        self._r.value = 'Error'
        self._R.value = 'Error'

    def __itemChanged(self, index):
            if self._combobox.value == 'Triangle':
                self._x4.readonly = True
                self._y4.readonly = True
            else:
                self._x4.readonly = False
                self._y4.readonly = False


def write_docx(figure, coords, s, r, R):
    document = Document('geom.docx')
    document.add_paragraph(figure)
    document.add_paragraph('coordinates: ')
    for coord in coords:
        document.add_paragraph('(' + str(coord[0]) + ', ' + str(coord[1]) + ')')
    document.add_paragraph('Площадь:' + str(round(s, 5)))
    document.add_paragraph('Радиус вписанной окружности:' + str(round(r, 5)) if r else 'вписать окружность невозможно')
    document.add_paragraph('Радиус описанной окружности:' + str(round(R, 5)) if R else 'описать окружность невозможно')
    document.save('geom.docx')


start_app(Geom_figures)
